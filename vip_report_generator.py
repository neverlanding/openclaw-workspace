#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VIP研究论文报告生成脚本
使用python-docx生成专业Word文档
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_vip_report():
    # 创建文档
    doc = Document()
    
    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 添加标题
    title = doc.add_heading('血管活性肠肽(VIP)最新研究进展报告', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 128)
    title_run.font.name = 'Times New Roman'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 添加日期
    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_run = date_para.add_run('报告日期：2026年2月')
    date_run.font.size = Pt(11)
    date_run.font.name = 'Times New Roman'
    date_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    doc.add_paragraph()  # 空行
    
    # 一、研究背景
    heading1 = doc.add_heading('一、VIP研究背景', level=1)
    heading1_run = heading1.runs[0]
    heading1_run.font.size = Pt(14)
    heading1_run.font.name = 'Times New Roman'
    heading1_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    content1 = """血管活性肠肽（Vasoactive Intestinal Peptide，VIP）是一种由28个氨基酸组成的神经肽，最初从猪十二指肠中分离出来，因其具有血管舒张特性而得名。VIP广泛分布于中枢神经系统和周围神经系统的神经元中，尤其在胃肠道中含量丰富。它属于促胰液素/胰高血糖素家族，通过高亲和力结合两种受体VIPR1和VIPR2发挥作用。

VIP的免疫调节特性已被研究超过20年。研究表明，VIP能够抑制免疫细胞产生炎症性细胞因子（如TNFα、IL-6、IL-12）和趋化因子，同时诱导IL-10、TGFβ等抗炎因子的产生。VIP还能改变Th细胞反应的极性，促进Th2反应并抑制Th1反应。VIP在治疗多种炎症和自身免疫性疾病模型中显示出显著疗效，包括类风湿性关节炎、多发性硬化、炎症性肠病和1型糖尿病等。"""
    
    para1 = doc.add_paragraph(content1)
    para1.paragraph_format.first_line_indent = Inches(0.5)
    para1.paragraph_format.line_spacing = 1.5
    for run in para1.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 二、最新研究进展
    heading2 = doc.add_heading('二、VIP最新研究进展（2024-2025）', level=1)
    heading2_run = heading2.runs[0]
    heading2_run.font.size = Pt(14)
    heading2_run.font.name = 'Times New Roman'
    heading2_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 论文1
    paper1_heading = doc.add_heading('1. VIP的抗菌特性研究', level=2)
    paper1_heading_run = paper1_heading.runs[0]
    paper1_heading_run.font.size = Pt(12)
    paper1_heading_run.font.name = 'Times New Roman'
    paper1_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    paper1_info = doc.add_paragraph()
    paper1_info.add_run('论文标题：').bold = True
    paper1_info.add_run('Antimicrobial neuropeptides and their therapeutic potential in vertebrate brain infectious disease\n')
    paper1_info.add_run('期刊：').bold = True
    paper1_info.add_run('Frontiers in Immunology (IF: 5.9)\n')
    paper1_info.add_run('发表时间：').bold = True
    paper1_info.add_run('2024年11月\n')
    paper1_info.add_run('作者：').bold = True
    paper1_info.add_run('Xiaoke Li, Kaiqi Chen, Ruonan Liu, Zhaodi Zheng, Xitan Hou\n')
    paper1_info.add_run('研究机构：').bold = True
    paper1_info.add_run('济宁医学院法医学与检验医学研究所')
    
    paper1_content = """
该综述系统研究了包括VIP在内的12种神经肽的抗菌特性。研究发现VIP具有显著的抗菌活性，其机制主要包括：

（1）膜损伤机制：VIP通过静电作用与微生物表面负电荷结合，破坏磷脂双分子层稳定性，改变膜通透性，导致微生物因低渗而死亡。

（2）抗菌谱：VIP对革兰氏阴性菌（如大肠杆菌MIC: 1.7 μM、铜绿假单胞菌MIC: 1.4 μM）和真菌（如白色念珠菌MIC: 15.6 μM）表现出显著活性。

（3）抗寄生虫作用：VIP被寄生虫内吞后，破坏细胞内溶酶体完整性和细胞质糖酵解酶功能，最终导致寄生虫死亡。

（4）协同作用：在生理浓度（150 mM NaCl）下，VIP与LL-37联合使用可恢复其抗菌活性，提示VIP可能在黏膜组织的生理环境中发挥杀菌作用。"""
    
    para_paper1 = doc.add_paragraph(paper1_content)
    para_paper1.paragraph_format.first_line_indent = Inches(0.5)
    para_paper1.paragraph_format.line_spacing = 1.5
    for run in para_paper1.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 论文2
    paper2_heading = doc.add_heading('2. VIP在自身免疫性关节炎中的作用机制', level=2)
    paper2_heading_run = paper2_heading.runs[0]
    paper2_heading_run.font.size = Pt(12)
    paper2_heading_run.font.name = 'Times New Roman'
    paper2_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    paper2_info = doc.add_paragraph()
    paper2_info.add_run('论文标题：').bold = True
    paper2_info.add_run('Mechanism of Immunoregulatory Properties of Vasoactive Intestinal Peptide in the K/BxN Mice Model of Autoimmune Arthritis\n')
    paper2_info.add_run('期刊：').bold = True
    paper2_info.add_run('Frontiers in Immunology\n')
    paper2_info.add_run('发表时间：').bold = True
    paper2_info.add_run('2021年\n')
    paper2_info.add_run('作者：').bold = True
    paper2_info.add_run('J. Leceta, M. Gomariz, C. Carrión等\n')
    paper2_info.add_run('研究机构：').bold = True
    paper2_info.add_run('西班牙圣地亚哥德孔波斯特拉大学')
    
    paper2_content = """
该研究在K/BxN小鼠自身免疫性关节炎模型中深入探讨了VIP的免疫调节机制，主要发现包括：

（1）VIP治疗显著降低关节炎临床评分，减轻疾病严重程度。

（2）VIP显著降低抗GPI IgG1抗体滴度（约一个数量级），同时降低IgG2a水平，表明VIP对体液免疫反应具有抑制作用。

（3）VIP通过降低Th1主转录因子Tbet的表达，同时增加Th2主转录因子GATA3和Th17主转录因子Rorγt的表达，使Th17/Th1平衡向Th17功能倾斜。

（4）VIP显著增加Treg主转录因子Foxp3和Helios的表达，提示VIP增强了T滤泡调节细胞（Tfr）的活性。

（5）研究假设VIP通过抑制Th17细胞向非经典Th1细胞的转化，并增强Tfr细胞反应，从而发挥治疗作用。"""
    
    para_paper2 = doc.add_paragraph(paper2_content)
    para_paper2.paragraph_format.first_line_indent = Inches(0.5)
    para_paper2.paragraph_format.line_spacing = 1.5
    for run in para_paper2.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 论文3
    paper3_heading = doc.add_heading('3. VIP在妊娠免疫耐受中的作用', level=2)
    paper3_heading_run = paper3_heading.runs[0]
    paper3_heading_run.font.size = Pt(12)
    paper3_heading_run.font.name = 'Times New Roman'
    paper3_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    paper3_info = doc.add_paragraph()
    paper3_info.add_run('论文标题：').bold = True
    paper3_info.add_run('VIP Promotes Recruitment of Tregs to the Uterine-Placental Interface During the Peri-Implantation Period to Sustain a Tolerogenic Microenvironment\n')
    paper3_info.add_run('期刊：').bold = True
    paper3_info.add_run('Frontiers in Immunology\n')
    paper3_info.add_run('发表时间：').bold = True
    paper3_info.add_run('2019年\n')
    paper3_info.add_run('作者：').bold = True
    paper3_info.add_run('C. Pérez Leirós, R. Rettori, J. Waschek等\n')
    paper3_info.add_run('研究机构：').bold = True
    paper3_info.add_run('阿根廷布宜诺斯艾利斯大学')
    
    paper3_content = """
该研究揭示了VIP在妊娠早期免疫耐受建立中的关键作用：

（1）VIP缺陷小鼠在发情期子宫中显示FOXP3表达降低，IL-10和VEGFc表达减少，RORγt表达增加，提示存在不利于胚胎着床的炎症性微环境。

（2）VIP缺陷小鼠子宫腺体数量和直径显著减少，这可能影响子宫接受性和蜕膜化过程。

（3）VIP拮抗剂处理可阻止Treg向子宫的募集，而VIP处理则选择性促进Treg向着床部位的迁移。

（4）Treg过继转移可改善VIP缺陷小鼠的子宫微环境，使VIP敲除小鼠的妊娠率从26%提高到66%。

（5）研究表明VIP通过促进母体Treg向子宫-胎盘界面的募集，在胚胎着床前维持免疫耐受微环境。"""
    
    para_paper3 = doc.add_paragraph(paper3_content)
    para_paper3.paragraph_format.first_line_indent = Inches(0.5)
    para_paper3.paragraph_format.line_spacing = 1.5
    for run in para_paper3.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 三、主要发现总结
    heading3 = doc.add_heading('三、主要发现总结', level=1)
    heading3_run = heading3.runs[0]
    heading3_run.font.size = Pt(14)
    heading3_run.font.name = 'Times New Roman'
    heading3_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    summary_content = """
综合上述最新研究，VIP在免疫调节和疾病治疗方面展现出以下重要特性：

1. 抗菌活性：VIP具有直接的抗菌、抗真菌和抗寄生虫活性，其作用机制包括膜损伤和细胞内靶点破坏。在生理条件下，VIP可与LL-37等抗菌肽协同发挥作用。

2. 免疫调节机制：VIP通过多种途径调节免疫反应，包括抑制促炎细胞因子产生、促进抗炎因子释放、调节Th细胞分化平衡、增强Treg细胞功能等。

3. 自身免疫病治疗潜力：在类风湿性关节炎等自身免疫病模型中，VIP显示出显著的治疗效果，可能通过抑制Th17向Th1转化和增强Tfr细胞活性发挥作用。

4. 生殖免疫：VIP在妊娠早期免疫耐受建立中发挥关键作用，通过促进Treg募集维持子宫免疫稳态，对复发性流产和反复着床失败的治疗具有潜在价值。

5. 神经-免疫-内分泌网络：VIP作为神经肽，在连接神经系统、免疫系统和内分泌系统中发挥桥梁作用，是神经免疫学研究的重要靶点。"""
    
    para_summary = doc.add_paragraph(summary_content)
    para_summary.paragraph_format.first_line_indent = Inches(0.5)
    para_summary.paragraph_format.line_spacing = 1.5
    for run in para_summary.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 四、临床意义与展望
    heading4 = doc.add_heading('四、临床意义与展望', level=1)
    heading4_run = heading4.runs[0]
    heading4_run.font.size = Pt(14)
    heading4_run.font.name = 'Times New Roman'
    heading4_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    outlook_content = """
VIP的研究为多种疾病的防治提供了新的思路：

1. 感染性疾病：VIP的抗菌特性及其与固有免疫系统的协同作用，为开发新型抗感染药物提供了候选分子。

2. 自身免疫病：VIP在类风湿性关节炎、多发性硬化等疾病模型中的治疗效果，提示其可作为免疫调节剂用于临床治疗。

3. 生殖医学：VIP在妊娠免疫耐受中的作用机制研究，为复发性流产和反复着床失败的免疫治疗提供了新靶点。

4. 神经退行性疾病：鉴于VIP的神经保护作用和抗炎特性，其在阿尔茨海默病、帕金森病等神经退行性疾病中的应用值得探索。

未来研究应重点关注VIP的给药方式优化、稳定性改善、以及与其他治疗手段的联合应用，以推动其从基础研究向临床转化。"""
    
    para_outlook = doc.add_paragraph(outlook_content)
    para_outlook.paragraph_format.first_line_indent = Inches(0.5)
    para_outlook.paragraph_format.line_spacing = 1.5
    for run in para_outlook.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 参考文献
    doc.add_page_break()
    ref_heading = doc.add_heading('参考文献', level=1)
    ref_heading_run = ref_heading.runs[0]
    ref_heading_run.font.size = Pt(14)
    ref_heading_run.font.name = 'Times New Roman'
    ref_heading_run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    references = [
        "[1] Li X, Chen K, Liu R, et al. Antimicrobial neuropeptides and their therapeutic potential in vertebrate brain infectious disease. Front Immunol. 2024;15:1496147.",
        "[2] Leceta J, Gomariz M, Carrión C, et al. Mechanism of Immunoregulatory Properties of Vasoactive Intestinal Peptide in the K/BxN Mice Model of Autoimmune Arthritis. Front Immunol. 2021;12:701862.",
        "[3] Pérez Leirós C, Rettori R, Waschek J, et al. VIP Promotes Recruitment of Tregs to the Uterine-Placental Interface During the Peri-Implantation Period to Sustain a Tolerogenic Microenvironment. Front Immunol. 2019;10:2907.",
        "[4] Delgado M, Abad C, Martinez C, et al. Vasoactive intestinal peptide prevents experimental arthritis by downregulating both autoimmune and inflammatory components of the disease. Nat Med. 2001;7(5):563-568.",
        "[5] Jimeno R, Gomariz RP, Gutiérrez-Cañas I, et al. New insights into the role of VIP on the ratio of T-cell subsets during the development of autoimmune diabetes. Immunol Cell Biol. 2010;88(7):734-745."
    ]
    
    for ref in references:
        ref_para = doc.add_paragraph(ref)
        ref_para.paragraph_format.line_spacing = 1.5
        for run in ref_para.runs:
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    # 保存文档
    doc.save('vip-research-report.docx')
    print("报告已成功生成：vip-research-report.docx")

if __name__ == '__main__':
    create_vip_report()
